from pathlib import Path
from typing import Optional
import tempfile
import logging

logger = logging.getLogger(__name__)

def extract_text_by_type(response, file_ext: str, filename: str) -> str:
    """
    根据文件类型提取文本
    """
    try:
        file_ext = file_ext.lower()

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(
            suffix=file_ext,
            delete=False,
        ) as temp_file:
            temp_file.write(response.read())
            temp_file_path = temp_file.name
            
        try:
            if file_ext == '.pdf':
                text = _extract_pdf_file(temp_file_path)
            elif file_ext == '.docx':
                text = _extract_docx_file(temp_file_path)
            elif file_ext in ['.txt', '.md']:
                text = _extract_text_file(temp_file_path)
            else:
                text = _extract_file_auto(temp_file_path)
            
            # 验证   
            _validate_extracted_text(text, filename)
            
            logger.info(f"Text extraction completed for file {filename}")
            
            return text
        
        finally:
            Path(temp_file_path).unlink(missing_ok=True)
            
    except Exception as e:
        logger.error(f"Text extraction error for file {filename}: {str(e)}", exc_info=True)
        raise
    
def _extract_pdf_file(file_path: str) -> str:
    """
    优化使用pdfplumber提取PDF文本
    """
    try:
        import pdfplumber     
        
        text = ""  
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text() or ""
                
                # 提取表格
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        text += "\n[Table]\n"
                        for row in table:
                            text += " | ".join(str(cell or "") for cell in row) + "\n"
                        text += "[/Table]\n"
                        
        return text
    
    except ImportError:
        logger.error("pdfplumber not installed, falling back to PyMuPDF")
        return _extract_pdf_fallback(file_path)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {str(e)}, trying fallback")
        return _extract_pdf_fallback(file_path)
       
    
def _extract_pdf_fallback(file_path: str) -> str:
    """PDF 提取降级方案 - 使用 PyMuPDF"""
    try:
        import fitz
        
        text = "" 
        doc = fitz.open(file_path) 

        for page_num in range(len(doc)):
            page = doc[page_num]
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.get_text()
        
        doc.close()
        return text

    except ImportError:
        logger.error("Neither pdfplumber nor PyMuPDF installed")
        raise
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}", exc_info=True)
        raise

def _extract_docx_file(file_path: str) -> str:
    """优化使用unstructured提取DOCX文本"""
    try:
        from unstructured.partition.docx import partition_docx
        
        elements = partition_docx(file_path)
        text = "\n".join([str(ele) for ele in elements])
        
        return text

    except ImportError:
        logger.warning("unstructured not installed, using python-docx")
        return _extract_docx_native(file_path)
    except Exception as e:
        logger.warning(f"unstructured failed: {str(e)}, trying native")
        return _extract_docx_native(file_path)
    
def _extract_docx_native(file_path: str) -> str:
    """DOCX 提取降级方案 - 使用 python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""

        for para in doc.paragraphs:
            text += para.text + "\n"
            
        for table in doc.tables:
            text += "[Table]\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
            text += "[/Table]\n"
            
        return 
    
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}", exc_info=True)
        raise
    
def _extract_text_file(file_path: str) -> str:
    """文本文件提取文本"""
    try:
        # 尝试使用不同的编码解码
        for encoding in ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # 容错处理
        with open(file_path, "r", encoding="utf-8", error="replace") as f:
            return f.read()

    except Exception as e:
        logger.error(f"Text file extraction error: {str(e)}", exc_info=True)
        raise
    
def _extract_file_auto(file_path: str) -> str:
    """自动识别文件类型并提取文本"""
    try:
        from unstructured.partition.auto import partition
        
        elements = partition(file_path)
        text = "\n".join([str(ele) for ele in elements])
        
        return text
    
    except ImportError:
        logger.warning("unstructured not available, reading as text")
        return _extract_text_file(file_path)
    except Exception as e:
        logger.error(f"Auto extraction error: {str(e)}", exc_info=True)
        return _extract_text_file(file_path)
    
def _validate_extracted_text(text: str, filename: str):
    """验证提取的文本"""
    if not text or not text.strip():
        raise ValueError(f"Empty text extracted from {filename}")

    # 最小长度检查
    if len(text.strip()) < 10:
        logger.warning(f"Very short text extracted from {filename}")
    
    # 最大长度检查（防止异常文件）
    if len(text) > 50_000_000:  # 50MB
        raise ValueError(f"Text too large ({len(text)} chars) from {filename}")
    
    logger.info(f"Text validation passed and length is {len(text)}")  
    
    
    
# def _extract_pdf_text(response, filename: str) -> str:
#     """PDF 文件提取文本"""
#     try:
#         import PyPDF2
#         from io import BytesIO
        
#         pdf_bytes = BytesIO(response.read())
#         pdf_reader = PyPDF2.PdfReader(pdf_bytes)
        
#         text = ""
#         for page_num in range(len(pdf_reader.pages)):
#             try:
#                 page = pdf_reader.pages[page_num]
#                 text += f"\n--- Page {page_num + 1} ---\n"
#                 text += page.extract_text()
#             except Exception as e:
#                 logger.error(f"Error extracting text from page {page_num + 1}: {e}")
#                 continue
            
#         return text
#     except ImportError:
#         logger.error("PyPDF2 not installed")
#         raise
#     except Exception as e:
#         logger.error(f"PDF extraction error: {str(e)}", exc_info=True)
#         raise
    
# def _extract_docx_text(response) -> str:
#     """DOCX 文件提取文本"""
#     try:
#         from docx import Document
#         from io import BytesIO
        
#         doc_bytes = BytesIO(response.read())
#         doc = Document(doc_bytes)
        
#         def _extract_body_content(doc) -> str:
#             """遍历 document body 的所有块级元素"""
#             text_parts = []
            
#             # 遍历 document.element.body 的所有子元素
#             for block in doc.element.body:
#                 # 处理段落
#                 if isinstance(block, docx.oxml.text.paragraph.CT_P):
#                     para_text = _extract_paragraph_content(block)
#                     if para_text:
#                         text_parts.append(para_text)
                        
#                 # 处理表格
#                 elif isinstance(block, docx.oxml.text.table.CT_Tbl):
#                     table_text = _extract_table_with_structure(block, doc)
#                     if table_text:
#                         text_parts.append(table_text)
                        
#             return "\n".join(text_parts)
        
#         text = _extract_body_content(doc)
        
#         return text
        
#     except ImportError:
#         logger.error("python-docx not installed")
#         raise
#     except Exception as e:
#         logger.error(f"DOCX extraction error: {str(e)}", exc_info=True)
#         raise
    
# def _extract_paragraph_content(para_element) -> str:
#     """
#     提取段落内容：
#     - 普通文本
#     - 列表项标记
#     - 标题级别
#     - 超链接
#     """
#     from docx.omxl.ns import qn
    
 
#     style = para_element.get(qn('pStyle'))
#     pPr = para_element.get(qn('pPr'))
    
#     # 识别标题级别
#     if style is not None:
#         style_val = style.get(qn('val'))
#         if style_val and style_val.startswith('Heading'):
#             heading_level = int(style_val.replace('Heading', ''))
    
#     # 提取文本内容和格式       
#     text_runs = []
#     for run in para_element.findall(qn('r')):
#         text_elem = run.find(qn('t'))
#         if text_elem is not None and text_elem.text:
#             # 检查格式
#             rPr = run.find(qn('rPr'))
#             text = text_elem.text
        
#         # 处理粗体
#         if rPr is not None and rPr.find(qn('b')) is not None:
#             text = f"**{text}**"
        
#         # 处理斜体
#         if rPr is not None and rPr.find(qn('i')) is not None:
#             text = f"*{text}*"    
        
#         # 处理超链接
#         hyperlink = rPr.find(qn('hyperlink'))
#         if hyperlink is not None:
#             rId = hyperlink.get(qn('r:id'))
#             if rId:
#                 text = f"[{text}](link:{rId})"
                
#         text_runs.append(text)
        
#     text = "".join(text_runs)
        
#     # 添加标题标记
#     if heading_level:
#         return "#" * heading_level + " " + text
    
#     # 处理列表项
#     if pPr is not None:
#         numPr = pPr.find(qn('numPr'))
#         if numPr is not None:
#             ilvl = numPr.find(qn('ilvl'))
#             if ilvl is not None:
#                 level = int(ilvl.get(qn('val')))
#                 indent = " " * level
                
#                 # 检查是否是有序列表
#                 numId = numPr.find(qn('numId'))
#                 if numId is not None:
#                     return f"{indent}* {text}"
                
#     return text
        
# def _extract_table_with_structure(tbl_element, doc) -> str:
#     """提取表格内容，并保留结构"""    
#     from docx.table import Table
#     try:
#         table = Table(tbl_element, doc)
        
#         table_lines = []    

#         for row_idx, row in enumerate(table.rows):
#             row_cells = []
#             for cell in row.cells:
#                 cell_text = _extract_cell_content(cell)
#                 row_cells.append(cell_text)
                
#             row_line = " | ".join(row_cells)
#             table_lines.append(row_line)
            
#             if row_idx == 0:
#                 separator = " | ".join(["-"] * max(10, len(cell)) for cell in row_cells)
#                 table_lines.append(separator)
                
#         table_text = "\n[Table]\n" + "\n".join(table_lines) + "\n[/Table]\n"
        
#         return table_text
    
#     except Exception as e:
#         logger.warning(f"Failed to extract table structure: {str(e)}")
#         return _extract_table_fallback(tbl_element)
    
# def _extract_cell_content(cell) -> str:
#     cell_texts = []

#     for para in cell.paragraphs:
#         para_text = ""
#         for run in para.runs:
#             if run.text:
#                 para_text += run.text
#             if para_text.strip():
#                 cell_texts.append(para_text.strip())
                
#     for tbl in cell.tables:
#         nested_table = _extract_table_with_structure(tbl._element, cell._parent)
#         if nested_table:
#             cell_texts.append(f"[Nested Table]\n{nested_table}")
    
#     return "\n".join(cell_texts)

# def _extract_table_fallback(tbl_element) -> str:
#     """
#     表格提取失败时的降级方案
#     """
#     text_parts = []
#     text_parts.append("[Table Start]")
    
#     for tr in tbl_element.findall(".//tr"):
#         row_data = []
#         for tc in tr.findall(".//tc"):
#             cell_text = ""
#             for para in tc.findall(".//p"):
#                 for t in para.findall(".//t"):
#                     if t.text:
#                         cell_text += t.text
#             row_data.append(cell_text.strip())
        
#         if row_data:
#             text_parts.append(" | ".join(row_data))
    
#     text_parts.append("[Table End]")
#     return "\n".join(text_parts)
    
        

        